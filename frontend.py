from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
import numpy as np
from pathlib import Path
downloads_path = str(Path.home() / "Downloads")
import os
import io
import time

import streamlit as st
st.set_page_config(
    page_title = 'Pinnacle Accounting Cover Letter2',
    page_icon = ':white_check_mark:',
    layout = 'wide'
)
from streamlit_gsheets import GSheetsConnection
doc = Document("data/coverpage.docx") 
# Create a connection object.


if 'user_name' not in st.session_state:
    st.session_state['user_name'] = ''

if 'password' not in st.session_state:
    st.session_state['password'] = ''

if "word_doc_button_clicked" not in st.session_state:
    st.session_state.word_doc_button_clicked = False


login = None
conn = st.connection("gsheets", type=GSheetsConnection)

def password_checker(user_name,user_password):
    user_crediantials = conn.read(worksheet="Crediantials")
    user_crediantials_list = user_crediantials["username"].to_list()
    user_crediantials = user_crediantials.set_index("username")

    if  user_name in user_crediantials_list:
        # print("Correct Username")
        if user_crediantials["password"][user_name] == user_password:
            # login = True
            return True
        
        return "Password is incorrect"
    
    return "Username is incorrect"


def add_new_doc(df):
    for x in range(len(df)):
        # print(df["Content"][x])

        if df["Type"][x] == "new paragraph":
            paragraph  = doc.add_paragraph()

        elif df["Type"][x] == "paragraph":
            font = paragraph.add_run(f'''{df["Content"][x]}''').font
            font.name = df["fontname"][x]
            font.size = Pt(df["fontsize"][x])
            font.bold = df["fontbold"][x] 
            font.underline  = df["fontunderline"][x] 
            font.highlight_color  = df["fonthighlight_color"][x]  

        elif df["Type"][x] == "table": 
            table = doc.add_table(rows=0, cols=2)
            table.style = doc.styles['Table Grid']
            table.autofit = True
            table.allow_autofit = True

        elif df["Type"][x] == "add on to table": 
            row_cells = table.add_row().cells
            tableparagraph = row_cells[0].paragraphs[0]
            run = tableparagraph.add_run(f"{df['Content'][x]}")
            run.underline = df["fontunderline"][x] 
            run.font.size = Pt(df["fontsize"][x])
            run.font.name = df["fontname"][x]
            run.font.highlight_color = df["fonthighlight_color"][x]
            tableparagraph = row_cells[1].paragraphs[0]
            run = tableparagraph.add_run(f"{df['content2'][x]}")
            run.underline = df["fontunderline"][x] 
            run.font.size = Pt(df["fontsize"][x])
            run.font.name = df["fontname"][x]
            run.font.highlight_color = df["fonthighlight_color"][x]

        else:
            pass

    return doc




def sel_callback():
    st.session_state.col1 = st.session_state.sel
    st.session_state.col2 = st.session_state.sel
    st.session_state.col3 = st.session_state.sel
    st.session_state.col4 = st.session_state.sel
    st.session_state.col5 = st.session_state.sel
    st.session_state.col6 = st.session_state.sel
    st.session_state.col7 = st.session_state.sel
    st.session_state.col8 = st.session_state.sel
    st.session_state.col9 = st.session_state.sel
    st.session_state.col10 = st.session_state.sel
    st.session_state.col11 = st.session_state.sel
    st.session_state.col12 = st.session_state.sel
    st.session_state.col13 = st.session_state.sel
    st.session_state.col14 = st.session_state.sel
    st.session_state.col15 = st.session_state.sel
    st.session_state.col16 = st.session_state.sel
    st.session_state.col17 = st.session_state.sel
    st.session_state.col18 = st.session_state.sel
    st.session_state.col19 = st.session_state.sel
    st.session_state.col20 = st.session_state.sel
    st.session_state.col21 = st.session_state.sel
    st.session_state.col22 = st.session_state.sel
    st.session_state.col23 = st.session_state.sel
    st.session_state.col24 = st.session_state.sel
    st.session_state.col25 = st.session_state.sel


# if login == True:

st.header(":ballot_box_with_check: Select the content you would like to include in the Accounting Letter")
st.write("")
cola, colb,colc= st.columns(3)



with cola:
    Property_plant_equipment = st.checkbox("Property plant equipment", value=False, key="col1")
    Investment_properties = st.checkbox("Investment properties", value=False, key="col2")
    Fair_value = st.checkbox("Fair value", value=False, key="col3")
    Investment_in_joint_venture = st.checkbox("Investment in joint venture", value=False, key="col4")
    Investment_in_associates = st.checkbox("Investment in associates", value=False, key="col5")
    Investment_in_subsidiaries = st.checkbox("Investment in subsidiaries", value=False, key="col6")
    Intangible_assets = st.checkbox("Intangible assets", value=False, key="col7")
    Goodwill = st.checkbox("Goodwill", value=False, key="col8")
    st.write("")
    st.checkbox('Select All', key='sel', on_change=sel_callback)

with colb:
    Right_of_use = st.checkbox("Right of use", value=False, key="col9")
    Inventory = st.checkbox("Inventory", value=False, key="col10")
    Cash_and_bank_balances = st.checkbox("Cash and bank balances", value=False, key="col11")
    Trade_and_other_receivables = st.checkbox("Trade and other receivables", value=False, key="col12")
    Trade_and_other_payables = st.checkbox("Trade and other payables", value=False, key="col13")
    Borrowings = st.checkbox("Borrowings", value=False, key="col14")
    Amount_due = st.checkbox("Amount due", value=False, key="col15")
    Revenue_recognition = st.checkbox("Revenue recognition", value=False, key="col16")
    Gross_profit_margin = st.checkbox("Gross profit margin", value=False, key="col17")

with colc:
    Voluntary_cpf_contributions = st.checkbox("Voluntary cpf contributions", value=False, key="col18")
    Government_grant = st.checkbox("Government grant", value=False, key="col19")
    Foreign_exchange = st.checkbox("Foreign exchange", value=False, key="col20")
    Small_value_assets = st.checkbox("Small value assets", value=False, key="col21")
    Presentation_currency = st.checkbox("Presentation currency", value=False, key="col22")
    Expenses_recognition = st.checkbox("Expenses recognition", value=False, key="col23")
    Gst_registration = st.checkbox("Gst registration", value=False, key="col24")
    Representations_from_the_company = st.checkbox("Representations from the company", value=False, key="col25")   

mehmehlist = [Property_plant_equipment,Investment_properties, Fair_value, Investment_in_joint_venture,
Investment_in_associates,Investment_in_subsidiaries, Intangible_assets,Goodwill,Right_of_use,
Inventory, Cash_and_bank_balances, Trade_and_other_receivables,
Trade_and_other_payables,Borrowings,Amount_due,Revenue_recognition,Gross_profit_margin,
Voluntary_cpf_contributions,Government_grant,Foreign_exchange,Small_value_assets, Presentation_currency,
Expenses_recognition,Gst_registration,Representations_from_the_company]

list_of_sheets = ["PROPERTY, PLANT AND EQUIPMENT", "INVESTMENT PROPERTIES", "FAIR VALUE", 
"INVESTMENT IN JOINT VENTURE", "INVESTMENT IN ASSOCIATES", "INVESTMENT IN SUBSIDIARIES",
"INTANGIBLE ASSETS", "GOODWILL", "RIGHT OF USE ASSETS & LEASE LIABILITIES", "INVENTORIES",
"CASH AND BANK BALANCES", "TRADE AND OTHER RECEIVABLES", "TRADE AND OTHER PAYABLES", "BORROWINGS",
"AMOUNT DUE FROM/ TO SHAREHOLDERS/ DIRECTORS", "REVENUE RECOGNITION, CONTRACT ASSETS AND CONTRACT LIABILITIES",
"GROSS PROFIT MARGIN", "VOLUNTARY CPF CONTRIBUTIONS", "GOVERNMENT GRANT – CAPITAL", "FOREIGN EXCHANGE", 
"SMALL VALUE ASSETS", "PRESENTATION CURRENCY", "EXPENSES RECOGNITION", "GST REGISTRATION", "REPRESENATATIONS FROM THE COMPANY"]



def checker():
    doc = Document("data/coverpage.docx") 
    tick = 0
    for x in range(len(mehmehlist)):
        if mehmehlist[x] == True:
            # print(mehmehlist[x])
            df = conn.read(worksheet=list_of_sheets[x])
            df = df.replace(np.nan,'', regex=True)
            df["fonthighlight_color"] = df["fonthighlight_color"].replace("None",None,regex=True)
            df["fonthighlight_color"] = df["fonthighlight_color"].replace("turquoise",WD_COLOR_INDEX.TURQUOISE,regex=True)
            df["fonthighlight_color"] = df["fonthighlight_color"].replace("yellow",WD_COLOR_INDEX.YELLOW,regex=True)
            doc = add_new_doc(df)
            tick += 1
            if tick == 10:
                time.sleep(15)
            elif tick == 20:
                time.sleep(15)

    return doc



def Word_Document_Created():
    st.session_state.word_doc_button_clicked = True

username = st.text_input("Username", key="username")
password = st.text_input("Password", key="password", type="password")
docname = st.text_input(label="Document Name",placeholder = "Insert Word Document File Name Here Before Clicking Create")

if ( 
    st.button("Create Word Document")
):
    
    check = password_checker(username,password)
    if check == True:
    
        st.subheader("After loading is done, press the button below to download. File will be in your download folder")
        
        doc = checker()
        # doc.save(f'{downloads_path}/{docname}.docx')
        bio = io.BytesIO()
        doc.save(bio)
        
        if doc:
                st.download_button(
                    label="Click here to download",
                    data=bio.getvalue(),
                    file_name=f"{docname}.docx",
                    mime="docx"
                )
    else:
        st.error(check)
        st.session_state.word_doc_button_clicked = False


# else:
#     user_name = st.text_input('User email', placeholder = 'username@email.com')
#     user_password =  st.text_input('Password', placeholder = '12345678',type="password")
#     login = st.button("Login", type="primary")
#     user_crediantials = conn.read(worksheet="Crediantials")
#     st.session_state['user_name'] = user_name
#     st.session_state['password'] = user_password

















# for x in mehmeh:
    

# doc.save(f'{downloads_path}/excelonlinetest14.docx')

            

